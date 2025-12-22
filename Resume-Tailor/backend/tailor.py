import openai
import os
import json
import shutil
from docx import Document
from typing import List, Dict

# Ensure API key is set
# openai.api_key = os.environ.get("OPENAI_API_KEY")

from pydantic import BaseModel

class Edit(BaseModel):
    target_text: str
    new_content: str
    action: str
    rationale: str

class SectionAnalysis(BaseModel):
    section_name: str
    gaps: List[str]
    suggestions: List[str]
    edits: List[Edit]

class AnalysisResult(BaseModel):
    sections: List[SectionAnalysis]
    initial_score: int
    projected_score: int

def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    return '\n'.join(full_text)

def safe_replace_text(paragraph, target: str, replacement: str):
    """
    Attempts to replace 'target' with 'replacement' in the paragraph while preserving formatting.
    Strategy:
    1. Check if 'target' exists in a single run. If so, replace it there.
    2. If not, fallback to paragraph.text replacement but attempt to copy font from the first run of the target.
    """
    if target not in paragraph.text:
        return

    # 1. Try single run replacement
    for run in paragraph.runs:
        if target in run.text:
            run.text = run.text.replace(target, replacement)
            return

    # 2. Fallback: Replace whole paragraph text but try to keep style
    # We will clear the paragraph and add a new run with the replacement text,
    # applying the font/style from the first run (if any).
    # NOTE: This replaces the ENTIRE paragraph content if we do this, which might be too aggressive 
    # if the target is just a SUBSTRING. 
    
    # Better Fallback for Substrings spanning runs:
    # Just do a naive replace on text. This destroys run boundaries for that paragraph.
    # To mitigate "lost font", we capture the first run's font.
    
    style_run = paragraph.runs[0] if paragraph.runs else None
    font_name = style_run.font.name if style_run else None
    font_size = style_run.font.size if style_run else None
    bold = style_run.bold if style_run else None
    italic = style_run.italic if style_run else None
    color = style_run.font.color.rgb if style_run and style_run.font.color else None

    # This operation clears existing runs and creates new ones usually? 
    # Actually python-docx `para.text = ...` preserves the PARAGRAPH style but resets character formatting.
    paragraph.text = paragraph.text.replace(target, replacement)
    
    # Re-apply font to all runs (usually just one now)
    for run in paragraph.runs:
        if font_name: run.font.name = font_name
        if font_size: run.font.size = font_size
        if bold is not None: run.bold = bold
        if italic is not None: run.italic = italic
        if color: run.font.color.rgb = color


def apply_edits_to_docx(docx_path: str, edits: List[Dict[str, str]], output_path: str):
    # Copy original to output
    shutil.copy2(docx_path, output_path)
    doc = Document(output_path)
    
    for edit in edits:
        if isinstance(edit, (dict, object)): 
             # Handle both pydantic and dict
             if isinstance(edit, dict):
                action = edit.get("action")
                target = edit.get("target_text")
                content = edit.get("new_content")
             else:
                action = edit.action
                target = edit.target_text
                content = edit.new_content
        
        if not target or not content:
            continue
            
        for para in doc.paragraphs:
            if target in para.text:
                if action == "replace":
                    safe_replace_text(para, target, content)
                elif action == "append":
                     para.add_run(" " + content)
                break 

    doc.save(output_path)

def analyze_gaps(docx_path: str, job_description: str) -> Dict:
    resume_text = extract_text_from_docx(docx_path)
    
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    prompt = f"""
    You are an expert Resume Strategist.
    
    GOAL: 
    Tailor the resume to significantly increase the chances of being shortlisted.
    Analyze the resume SECTION BY SECTION.
    
    CONSTRAINTS:
    1. STRICTLY PRESERVE the original file structure. 
    2. Suggest specific text replacements.
    
    SCORING:
    - Assess the initial resume against the JD (0-100).
    - Estimate the projected score after your edits (0-100).
    
    STRATEGIES (Section-Specific):
    
    1. **Professional Summary**:
       - Ensure it matches the Job Title in the JD.
       - Highlight top 3 achievements relevant to the JD.
       - Use keywords from the JD (soft & hard skills).
    
    2. **Experience**:
       - Quantify results using the STAR method (Situation, Task, Action, Result) where possible.
       - Use strong action verbs (e.g., "Spearheaded," "Optimized," "Architected").
       - Explicitly integrate key technical and functional keywords from the JD.
       - Remove irrelevant responsibilities that dilute the impact.
    
    3. **Projects**:
       - Highlight technical problem-solving.
       - Mention specific technologies used (matching JD if applicable).
       - Focus on the *impact* of the project.
       
    4. **Skills**:
       - Prioritize hard skills found in the JD.
       - Group skills logically if they aren't already.
       
    5. **Education**:
       - Keep it concise.
    
    OUTPUT FORMAT:
    Return a JSON object:
    {{
        "initial_score": <int>,
        "projected_score": <int>,
        "sections": [
            {{
                "section_name": "<name>",
                "original_text": "<full original text of this section>",
                "gaps": ["<gap1>", ...],
                "suggestions": ["<high-level advice 1>", "<advice 2>", ...],
                "edits": [
                    {{
                        "target_text": "<exact match>",
                        "new_content": "<replacement>",
                        "action": "replace",
                        "rationale": "<reason>"
                    }}
                ]
            }}
        ]
    }}
    
    Job Description:
    {job_description}
    
    Original Resume Content:
    {resume_text}
    """
    
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        response_format={ "type": "json_object" }
    )
    
    try:
        result = json.loads(response.choices[0].message.content)
    except json.JSONDecodeError:
        print("Failed to decode JSON from LLM")
        result = {"sections": []}

    # 2. Separate Robust Scoring Step
    try:
        # Summarize proposed changes for the scorer
        changes_summary = []
        if "sections" in result:
            for section in result["sections"]:
                name = section.get("section_name", "Unknown")
                gaps = section.get("gaps", [])
                suggestions = section.get("suggestions", [])
                edits = section.get("edits", [])
                changes_summary.append(f"Section {name}: Found {len(gaps)} gaps. {len(suggestions)} suggestions. Suggested {len(edits)} edits.")
                if suggestions:
                    changes_summary.append(f" - Advice: {'; '.join(suggestions[:3])}")
        
        changes_text = "\n".join(changes_summary)
        scores = calculate_scores(resume_text, job_description, changes_text)
        result["initial_score"] = scores.get("initial_score", 0)
        result["projected_score"] = scores.get("projected_score", 0)
        result["score_reasoning"] = scores.get("reasoning", "")
        
    except Exception as e:
        print(f"Scoring failed: {e}")
        result["initial_score"] = 0
        result["projected_score"] = 0

    return result

def calculate_scores(resume_text: str, job_description: str, changes_summary: str) -> Dict:
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    prompt = f"""
    You are a Hiring Manager and ATS Specialist.
    
    JOB DESCRIPTION:
    {job_description[:2000]}...
    
    CANDIDATE RESUME CONTENT:
    {resume_text[:3000]}...
    
    PROPOSED IMPROVEMENTS TO RESUME:
    {changes_summary}
    
    TASK:
    1. Evaluate the *original* resume's match to the JD on a scale of 0-100 (ATS Score).
    2. Estimate the match score (0-100) assuming the proposed improvements are applied effectively.
    
    OUTPUT JSON:
    {{
        "initial_score": <int>,
        "projected_score": <int>,
        "reasoning": "<short explanation>"
    }}
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            response_format={ "type": "json_object" }
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"Error in calculate_scores: {e}")
        return {"initial_score": 0, "projected_score": 0}

def generate_tailored_resume(docx_path: str, sections: List[Dict]) -> str:
    # Flatten edits from all sections
    all_edits = []
    
    # Handle if sections is actually the AnalysisResult dict or list
    iterable_sections = sections if isinstance(sections, list) else sections.get("sections", [])
    
    for section in iterable_sections:
        # Handle dict vs object
        if isinstance(section, dict):
            if "edits" in section:
                all_edits.extend(section["edits"])
        else:
             all_edits.extend(section.edits)
            
    output_path = docx_path.replace(".docx", "_tailored.docx")
    apply_edits_to_docx(docx_path, all_edits, output_path)
    return output_path

