from pdf2docx import Converter
from docx2pdf import convert
import pythoncom
from docx import Document
from docx.oxml.ns import qn

def sanitize_docx_layout(docx_path: str):
    """
    Removes manual page breaks and section breaks to allow content to flow naturally.
    """
    doc = Document(docx_path)
    
    # 1. Remove 'page_break_before' from paragraphs
    for para in doc.paragraphs:
        if para.paragraph_format.page_break_before:
            para.paragraph_format.page_break_before = False
            
    # 2. Remove manual breaks (<w:br w:type="page">) in runs
    # This requires accessing the XML directly as python-docx doesn't fully expose breaks in runs easily
    for para in doc.paragraphs:
        for run in para.runs:
            if 'lastRenderedPageBreak' in run._element.xml:
                # This is just a marker, usually harmless but good to know
                pass
            
            # Look for <w:br> elements
            brs = run._element.findall(qn('w:br'))
            for br in brs:
                # Remove the break element
                run._element.remove(br)
                
    doc.save(docx_path)

def pdf_to_docx(pdf_path: str) -> str:
    docx_path = pdf_path.replace(".pdf", ".docx")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    
    # Sanitize immediately after conversion
    sanitize_docx_layout(docx_path)
    
    return docx_path

def docx_to_pdf(docx_path: str) -> str:
    # Initialize COM library for Windows
    pythoncom.CoInitialize()
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path

